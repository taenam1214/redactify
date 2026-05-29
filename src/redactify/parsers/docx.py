"""DOCX parser using python-docx."""

from pathlib import Path

from redactify.parsers.base import BaseParser, DocumentChunk, ParsedDocument


class DocxParser(BaseParser):
    """Parser for DOCX files using python-docx."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a DOCX file, extracting text from paragraphs and tables."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install redactify[docx]"
            )

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = Document(file_path)
        paragraphs = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs)
        chunk = DocumentChunk(text=full_text)

        return ParsedDocument(
            chunks=[chunk],
            source_path=file_path,
            file_type="docx",
        )

    def can_handle(self, file_path: Path) -> bool:
        return Path(file_path).suffix.lower() in {".docx", ".doc"}
