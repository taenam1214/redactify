"""Tests for the DOCX parser."""

import pytest

try:
    from redactify.parsers.docx import DocxParser
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestDocxParser:
    def setup_method(self):
        self.parser = DocxParser()

    def test_supported_extensions(self):
        from pathlib import Path
        assert self.parser.can_handle(Path("test.docx"))
        assert not self.parser.can_handle(Path("test.txt"))


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestDocxIntegration:
    def test_scan_docx_with_engine(self, tmp_path):
        """Create a simple DOCX with PII and scan it."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from redactify.core.engine import RedactionEngine

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Call us at 555-123-4567 for support.")
        doc.add_paragraph("Email: alice@company.com")
        doc.save(str(docx_path))

        engine = RedactionEngine(use_ner=False)
        report = engine.scan(docx_path)
        assert report.total_entities >= 2
